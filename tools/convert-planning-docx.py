"""Convert the 2026 planning DOCX archive into private Vault Markdown.

The mapping is intentionally explicit: this avoids silently putting a source
document in the wrong month and makes the conversion safe to rerun.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


MONTHS = {
    "一月": ("C2-1月计划", "01"),
    "二月": ("C3-2月计划", "02"),
    "三月": ("C4-3月计划", "03"),
    "四月": ("C5-4月计划", "04"),
    "五月": ("C6-5月计划", "05"),
    "六月": ("C7-6月计划", "06"),
    "七月": ("C8-7月计划", "07"),
}

TYPE_ORDER = {"日记": "D1", "规划": "D2", "日志": "D3"}
CONTENT_TYPES = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
}


@dataclass
class ConversionResult:
    source: Path
    target: Path
    paragraphs: int
    tables: int
    images: int
    source_characters: int
    body_characters: int


def yaml_quote(value: str) -> str:
    return '"' + value.replace("\\", "\\\\").replace('"', '\\"') + '"'


def markdown_cell(value: str) -> str:
    return value.replace("\\", "\\\\").replace("|", "\\|").replace("\r", "").replace("\n", "<br>")


def paragraph_text(paragraph: Paragraph) -> str:
    # Paragraph.text omits some hyperlink and tracked-insertion text. Reading
    # w:t nodes retains the visible text without accepting deleted revisions.
    parts: list[str] = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:t"):
            parts.append(node.text or "")
        elif node.tag == qn("w:tab"):
            parts.append("\t")
        elif node.tag in {qn("w:br"), qn("w:cr")}:
            parts.append("\n")
    return "".join(parts).strip()


def image_references(paragraph: Paragraph, document, asset_dir: Path, asset_link_dir: str,
                     image_names: dict[str, str]) -> list[str]:
    references: list[str] = []
    for blip in paragraph._p.iter(qn("a:blip")):
        relationship_id = blip.get(qn("r:embed"))
        if not relationship_id or relationship_id not in document.part.rels:
            continue
        relationship = document.part.rels[relationship_id]
        if "image" not in relationship.reltype:
            continue

        if relationship_id not in image_names:
            part = relationship.target_part
            extension = CONTENT_TYPES.get(part.content_type, Path(str(part.partname)).suffix or ".bin")
            filename = f"image-{len(image_names) + 1:02d}{extension.lower()}"
            asset_dir.mkdir(parents=True, exist_ok=True)
            (asset_dir / filename).write_bytes(part.blob)
            image_names[relationship_id] = filename

        references.append(f"![Word 内嵌图片]({asset_link_dir}/{image_names[relationship_id]})")
    return references


def paragraph_markdown(paragraph: Paragraph, document, asset_dir: Path, asset_link_dir: str,
                       image_names: dict[str, str]) -> tuple[list[str], str]:
    text = paragraph_text(paragraph)
    lines: list[str] = []
    if text:
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        heading_match = re.search(r"(?:Heading|标题)\s*([1-6])", style_name, re.IGNORECASE)
        if heading_match:
            text = f"{'#' * int(heading_match.group(1))} {text}"
        elif paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            text = f"- {text}"
        lines.append(text)
    lines.extend(image_references(paragraph, document, asset_dir, asset_link_dir, image_names))
    return lines, text


def table_markdown(table: Table, document, asset_dir: Path, asset_link_dir: str,
                   image_names: dict[str, str]) -> tuple[list[str], int]:
    rows: list[list[str]] = []
    image_lines: list[str] = []
    source_characters = 0
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            parts: list[str] = []
            for paragraph in cell.paragraphs:
                value = paragraph_text(paragraph)
                if value:
                    parts.append(value)
                    source_characters += len(re.sub(r"\s+", "", value))
                image_lines.extend(
                    image_references(paragraph, document, asset_dir, asset_link_dir, image_names)
                )
            cells.append(markdown_cell("\n".join(parts)))
        rows.append(cells)

    if not rows:
        return image_lines, source_characters
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    output = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    output.extend(image_lines)
    return output, source_characters


def classify_source(source: Path) -> tuple[str, str, str, str]:
    month = source.parent.name
    if month not in MONTHS:
        raise ValueError(f"不支持的月份目录：{month}")

    stem_without_prefix = re.sub(r"^C\d+-", "", source.stem)
    kind = next((candidate for candidate in TYPE_ORDER if candidate in stem_without_prefix), None)
    if kind is None and "工作日志" in stem_without_prefix:
        kind = "日志"
    if kind is None:
        raise ValueError(f"无法判断文件类型：{source.name}")

    target_month, month_number = MONTHS[month]
    target_name = f"{TYPE_ORDER[kind]}-{month}{kind}.md"
    title = f"2026年{stem_without_prefix}"
    return target_month, month_number, target_name, title


def convert(source: Path, vault_year_dir: Path) -> ConversionResult:
    target_month, month_number, target_name, title = classify_source(source)
    target = vault_year_dir / target_month / target_name
    target.parent.mkdir(parents=True, exist_ok=True)

    asset_dir = target.parent / f"{target.stem}-assets"
    asset_link_dir = f"{target.stem}-assets"
    document = Document(source)
    image_names: dict[str, str] = {}
    body_blocks: list[str] = []
    paragraph_count = 0
    table_count = 0
    source_characters = 0

    for element in document.element.body.iterchildren():
        if element.tag == qn("w:p"):
            paragraph = Paragraph(element, document)
            lines, visible_text = paragraph_markdown(
                paragraph, document, asset_dir, asset_link_dir, image_names
            )
            if visible_text:
                paragraph_count += 1
                source_characters += len(re.sub(r"\s+", "", paragraph_text(paragraph)))
            if lines:
                body_blocks.append("\n\n".join(lines))
        elif element.tag == qn("w:tbl"):
            table = Table(element, document)
            lines, table_characters = table_markdown(
                table, document, asset_dir, asset_link_dir, image_names
            )
            table_count += 1
            source_characters += table_characters
            if lines:
                body_blocks.append("\n".join(lines))

    body = "\n\n".join(block for block in body_blocks if block.strip()).strip()
    body_characters = len(re.sub(r"\s+", "", body))
    if source_characters == 0 or body_characters == 0:
        raise ValueError(f"转换结果为空，拒绝写入：{source}")

    kind = "日记" if "日记" in target_name else "月计划" if "规划" in target_name else "工作日志"
    frontmatter = "\n".join([
        "---",
        f"title: {yaml_quote(title)}",
        f"date: 2026-{month_number}-01 00:00:00 +0800",
        f"tags: [{yaml_quote('规划')}, {yaml_quote(kind)}, {yaml_quote(source.parent.name)}]",
        "published: false",
        "toc: true",
        "comments: false",
        "---",
        "",
    ])
    target.write_text(frontmatter + body + "\n", encoding="utf-8", newline="\n")

    return ConversionResult(
        source=source,
        target=target,
        paragraphs=paragraph_count,
        tables=table_count,
        images=len(image_names),
        source_characters=source_characters,
        body_characters=body_characters,
    )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--target", type=Path, required=True)
    args = parser.parse_args()

    sources = sorted(args.source.rglob("*.docx"), key=lambda path: str(path))
    if not sources:
        raise SystemExit("未找到 DOCX 源文件")

    results = [convert(source, args.target) for source in sources]
    print("source\ttarget\tparagraphs\ttables\timages\tsource_chars\tbody_chars")
    for result in results:
        print(
            f"{result.source.name}\t{result.target.relative_to(args.target)}\t"
            f"{result.paragraphs}\t{result.tables}\t{result.images}\t"
            f"{result.source_characters}\t{result.body_characters}"
        )
    print(f"TOTAL\t{len(results)} files")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
